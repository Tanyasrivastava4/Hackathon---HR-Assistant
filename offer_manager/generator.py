# offer_manager/generator.py
from docx import Document
from pathlib import Path
import os
from datetime import datetime

# -----------------------------
# Paths
# -----------------------------
TEMPLATE_PATH = Path(__file__).parent / "templates" / "offer_template.docx"
OUTPUT_DIR = Path(__file__).parent.parent / "data" / "reports" / "offers"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# -----------------------------
# Generate Offer Letter
# -----------------------------
def generate_offer_letter(candidate_name: str, position: str, salary: str, start_date: str, company_name: str = "Your Company"):
    """
    Reads the offer_template.docx file, replaces placeholders with actual values,
    and saves the personalized offer letter.
    """

    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Offer template not found at {TEMPLATE_PATH}. Please add a template.")

    # Load template
    doc = Document(str(TEMPLATE_PATH))

    # Prepare replacements
    replacements = {
        "{candidate_name}": candidate_name,
        "{position}": position,
        "{salary}": salary,
        "{start_date}": start_date,
        "{company_name}": company_name,
        "{date}": datetime.now().strftime("%d-%m-%Y")
    }

    # Replace placeholders in paragraphs
    for p in doc.paragraphs:
        for key, value in replacements.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Save the personalized offer letter
    safe_name = candidate_name.replace(" ", "_")
    output_file = OUTPUT_DIR / f"{safe_name}_offer.docx"
    doc.save(str(output_file))

    print(f"✅ Offer letter generated: {output_file}")
    return str(output_file)


# -----------------------------
# Example usage
# -----------------------------
if __name__ == "__main__":
    generate_offer_letter(
        candidate_name="Tanya Srivastava",
        position="Data Analyst Intern",
        salary="₹6,00,000",
        start_date="15-10-2025",
        company_name="Intileo Technologies"
    )





# offer_manager/generator.py
#from docx import Document
#from pathlib import Path
#import os

#TEMPLATE_PATH = Path(__file__).parent / "templates" / "offer_template.docx"
#OUTPUT_DIR = Path(__file__).parent.parent / "data" / "reports" / "offers"
#os.makedirs(OUTPUT_DIR, exist_ok=True)

#def generate_offer_letter(candidate_name: str, position: str, salary: str, start_date: str):
 #   if not TEMPLATE_PATH.exists():
  #      raise FileNotFoundError(f"Offer template not found at {TEMPLATE_PATH}. Please add a template.")

   # doc = Document(str(TEMPLATE_PATH))
    # Replace placeholders in paragraphs
   # for p in doc.paragraphs:
    #    inline = p.runs
     #   if inline:
      #      for i in range(len(inline)):
       #         text = inline[i].text
        #        text = text.replace("{candidate_name}", candidate_name)
         #       text = text.replace("{position}", position)
          #      text = text.replace("{salary}", salary)
           #     text = text.replace("{start_date}", start_date)
            #    inline[i].text = text

    # Replace placeholders in tables (if any)
  #  for table in doc.tables:
   #     for row in table.rows:
    #        for cell in row.cells:
     #           cell_text = cell.text
      #          cell_text = cell_text.replace("{candidate_name}", candidate_name)
       #         cell_text = cell_text.replace("{position}", position)
        #        cell_text = cell_text.replace("{salary}", salary)
         #       cell_text = cell_text.replace("{start_date}", start_date)
          #      cell.text = cell_text

  #  fname = f"{candidate_name.replace(' ', '_')}_offer.docx"
   # out_path = OUTPUT_DIR / fname
   # doc.save(str(out_path))
   # return str(out_path)

